"""
Test harness with file-based prompt (PDF converted from docx).
Task: Summarize vendor quotes and make a recommendation.
"""
import os
import sys
import subprocess
import traceback
from pathlib import Path
from dotenv import load_dotenv
load_dotenv()

from verif import RLHarness, ProviderConfig, Attachment, Prompt
from verif.providers.base import HistoryEntry
from verif.executor import SubprocessExecutor

# Path to the docx file (now at rubrics/ root, so example_tasks is sibling)
DOCX_PATH = Path(__file__).parent / "example_tasks" / "quotes_headlamp.docx"


def extract_docx_to_text(docx_path: Path) -> Path:
    """Extract text from DOCX and save to a .txt file for AI to read."""
    txt_path = docx_path.with_suffix(".txt")
    
    if txt_path.exists():
        print(f"Text file already exists: {txt_path}")
        return txt_path
    
    print(f"Extracting text from {docx_path.name}...")
    from docx import Document
    doc = Document(str(docx_path))
    
    # Extract all paragraphs and tables
    content = []
    for para in doc.paragraphs:
        if para.text.strip():
            content.append(para.text)
    
    # Extract tables
    for table in doc.tables:
        content.append("\n--- TABLE ---")
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            content.append(" | ".join(cells))
        content.append("--- END TABLE ---\n")
    
    text = "\n".join(content)
    txt_path.write_text(text)
    print(f"Created: {txt_path} ({len(text)} chars)")
    
    return txt_path


def print_event(entry: HistoryEntry):
    """Print harness events to console for visibility."""
    entry_type = entry.entry_type.upper()
    content = entry.content[:400] + "..." if len(entry.content) > 400 else entry.content
    # Color coding for different event types
    colors = {
        "SYSTEM": "\033[94m",    # Blue
        "MODEL": "\033[92m",     # Green
        "TOOL_CALL": "\033[93m", # Yellow
        "TOOL_RESPONSE": "\033[95m",  # Magenta
        "TOOL_ERROR": "\033[91m",     # Red
        "USER": "\033[96m",      # Cyan
    }
    color = colors.get(entry_type, "")
    reset = "\033[0m" if color else ""
    print(f"\n{color}[{entry_type}]{reset} {content}")
    sys.stdout.flush()


def main():
    # Extract text from DOCX (more reliable than PDF conversion)
    txt_path = extract_docx_to_text(DOCX_PATH)

    # Read preview for text file (first 100 lines)
    with open(txt_path, 'r', encoding='utf-8') as f:
        preview = ''.join(f.readlines()[:100])

    # Create attachment with file PATH + preview
    attachment = Attachment(
        content=str(txt_path),
        mime_type="text/plain",
        name="quotes_headlamp.txt",
        preview=preview,
    )
    
    # Task prompt
    TASK = f"""Attached is a text file with vendor quotations for Model I headlamp from three suppliers: Autolantic, Vendocrat, and Solimoto.

Analyze the quotes and produce a markdown summary containing:
1. A comparison table covering: unit price, tooling/NRE costs, lead time, FX exposure, and any other relevant terms
2. Key pricing differences between vendors (quantified where possible)
3. A recommendation on which vendor to select with justification

Save the output as vendor_comparison.md using execute_code.
"""

    # Configure artifacts directory for any generated files
    artifacts_dir = Path(__file__).parent / "artifacts"
    artifacts_dir.mkdir(exist_ok=True)
    
    # Create multimodal prompt with text + attachment
    prompt: Prompt = [
        TASK,
        attachment,
    ]
    
    # Configure harness - OpenAI with bash + code, NO web search
    config = ProviderConfig(name="openai", reasoning_effort="medium")
    harness = RLHarness(
        provider=config,
        enable_search=False,     # NO web search - not needed
        enable_bash=True,        # bash for file ops if needed
        enable_code=True,        # code execution to save files
        code_executor=SubprocessExecutor(str(artifacts_dir)),
        artifacts_dir=str(artifacts_dir),
        max_iterations=30,       # enough for this task
        on_event=print_event,    # print events to console
    )
    
    print(f"Running task with OpenAI + text file attachment")
    print(f"Text file: {txt_path}")
    print(f"Artifacts: {artifacts_dir}")
    print("-" * 50)
    
    try:
        result = harness.run_single(prompt)
        error = None
    except Exception as e:
        result = None
        error = f"**FATAL ERROR**\n```\n{traceback.format_exc()}\n```"
    
    # Save results to markdown
    output_file = Path(__file__).parent / "test_w_file_output_run3.md"
    with open(output_file, "w") as f:
        if error:
            f.write(f"# Error\n\n{error}\n\n---\n\n")
        
        if result:
            f.write(f"# Result\n\n")
            f.write(f"## Task\n{result.task}\n\n")
            f.write(f"## Answer\n{result.answer}\n\n")
            f.write(f"## Rubric\n{result.rubric}\n\n")
            f.write("---\n\n")
        
        # Add execution trace
        f.write(harness.get_history_markdown())
    
    print(f"\nResults saved to {output_file}")
    if error:
        print("RUN FAILED - see output file for details")
    else:
        print(f"\nAnswer preview:\n{result.answer[:500]}...")


if __name__ == "__main__":
    main()